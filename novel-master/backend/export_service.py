"""export_service.py — Novel Master Export Compiler
Compiles project files into unified export formats (txt, md, docx).
"""
import os
import sqlite3
from datetime import datetime

def compile_project_export(db, project_id, title, user_path, format_type='txt'):
    """Compile all text files in a project into a single export file.
    Returns (export_path, error_message).
    """
    files = db.execute(
        "SELECT file_path, display_name, word_count FROM manuscript_files "
        "WHERE project_id = ? AND mime_type IN (?, ?) ORDER BY sort_order, created_at",
        (project_id, 'text/plain', 'text/markdown')
    ).fetchall()

    if not files:
        return None, 'No text files found in project'

    export_dir = os.path.join(user_path, 'projects', str(project_id), 'exports')
    os.makedirs(export_dir, exist_ok=True)
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

    if format_type == 'txt':
        export_path = os.path.join(export_dir, f"{title}_{timestamp}.txt")
        with open(export_path, 'w', encoding='utf-8') as out:
            out.write(f"{title}\n")
            out.write("=" * 60 + "\n\n")
            for f in files:
                if os.path.exists(f['file_path']):
                    out.write(f"## {f['display_name']}\n\n")
                    with open(f['file_path'], 'r', encoding='utf-8') as src:
                        out.write(src.read())
                    out.write("\n\n")
        return export_path, None

    elif format_type == 'md':
        export_path = os.path.join(export_dir, f"{title}_{timestamp}.md")
        with open(export_path, 'w', encoding='utf-8') as out:
            out.write(f"# {title}\n\n")
            for f in files:
                if os.path.exists(f['file_path']):
                    out.write(f"## {f['display_name']}\n\n")
                    with open(f['file_path'], 'r', encoding='utf-8') as src:
                        out.write(src.read())
                    out.write("\n\n---\n\n")
        return export_path, None

    elif format_type == 'docx':
        try:
            from docx import Document
            from docx.shared import Pt
            export_path = os.path.join(export_dir, f"{title}_{timestamp}.docx")
            doc = Document()
            doc.add_heading(title, 0)
            for f in files:
                if os.path.exists(f['file_path']):
                    doc.add_heading(f['display_name'], level=1)
                    with open(f['file_path'], 'r', encoding='utf-8') as src:
                        content = src.read()
                    doc.add_paragraph(content)
                    doc.add_paragraph()
            doc.save(export_path)
            return export_path, None
        except ImportError:
            return None, 'python-docx not installed. Run: pip install python-docx'
        except Exception as e:
            return None, f'DOCX generation failed: {str(e)}'

    else:
        return None, f'Unsupported export format: {format_type}'
